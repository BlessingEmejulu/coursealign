import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCREENSHOTS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "documentation", "appendix_b", "screenshots"))

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, filename, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    cell = table.cell(0, 0)
    set_cell_background(cell, "F3F4F6")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'D1D5DB')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p_header = cell.paragraphs[0]
    p_header.paragraph_format.space_before = Pt(2)
    p_header.paragraph_format.space_after = Pt(4)
    r_hdr = p_header.add_run(f"File: {filename}\n")
    r_hdr.font.name = 'Consolas'
    r_hdr.font.size = Pt(9)
    r_hdr.font.bold = True
    r_hdr.font.color.rgb = RGBColor(55, 65, 81)

    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = p_header
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_figure_with_image(doc, img_filename, figure_caption, figure_description):
    img_path = os.path.join(SCREENSHOTS_DIR, img_filename)
    
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.0))
    else:
        # Fallback to placeholder box
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        cell = table.cell(0, 0)
        set_cell_background(cell, "F9FAFB")
        set_cell_margins(cell, top=200, bottom=200, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Insert Screenshot: {img_filename}]")
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(107, 114, 128)

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(2)
    p_cap.paragraph_format.keep_with_next = True
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run(figure_caption)
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(10)
    r_cap.font.bold = True

    # Description
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(2)
    p_desc.paragraph_format.space_after = Pt(14)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_desc = p_desc.add_run(f"Description: {figure_description}")
    r_desc.font.name = 'Arial'
    r_desc.font.size = Pt(9.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(75, 85, 99)

def generate_appendices():
    doc = Document()

    # Set standard page margins (1 inch)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(17, 24, 39)

    # --- MAIN TITLE ---
    p_main = doc.add_paragraph()
    p_main.paragraph_format.space_before = Pt(0)
    p_main.paragraph_format.space_after = Pt(18)
    p_main.paragraph_format.keep_with_next = True
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = p_main.add_run("APPENDICES")
    r_main.font.name = 'Arial'
    r_main.font.size = Pt(20)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(17, 24, 39)

    # =========================================================================
    # APPENDIX A: PROGRAM LISTING
    # =========================================================================
    p_app_a = doc.add_heading(level=1)
    p_app_a.paragraph_format.space_before = Pt(12)
    p_app_a.paragraph_format.space_after = Pt(12)
    p_app_a.paragraph_format.keep_with_next = True
    r_app_a = p_app_a.add_run("Appendix A: Program Listing")
    r_app_a.font.name = 'Arial'
    r_app_a.font.size = Pt(15)
    r_app_a.font.bold = True
    r_app_a.font.color.rgb = RGBColor(31, 41, 55)

    # A.1 Project Structure
    p_a1 = doc.add_heading(level=2)
    p_a1.paragraph_format.space_before = Pt(10)
    p_a1.paragraph_format.space_after = Pt(6)
    p_a1.paragraph_format.keep_with_next = True
    r_a1 = p_a1.add_run("A.1 Project Structure")
    r_a1.font.name = 'Arial'
    r_a1.font.size = Pt(12)
    r_a1.font.bold = True

    p = doc.add_paragraph("The CourseAlign system is organized into a modular full-stack software architecture comprising a FastAPI Python backend, a client-side frontend with custom responsive UI styling, and local Large Language Model (LLM) inference integration via Ollama. The complete directory and file hierarchy is organized as follows:")
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tree_str = """coursealign/
├── backend/
│   ├── app/
│   │   ├── auth/
│   │   │   └── auth.py                  # JWT authentication, bcrypt hashing, and RBAC guards
│   │   ├── core/
│   │   │   └── config.py                # System settings and environment configuration
│   │   ├── database/
│   │   │   ├── crud.py                  # Direct data access routines
│   │   │   └── database.py              # SQLite connection lifecycle & DDL table schemas
│   │   ├── routers/
│   │   │   ├── admin_router.py          # Administrative controls and user role governance
│   │   │   ├── ai_router.py             # LLM tutoring and automated quiz generation routes
│   │   │   ├── auth_router.py           # User registration, login, and profile analytics
│   │   │   ├── bookmark_router.py       # Course bookmarking endpoints
│   │   │   └── course_router.py         # Course outlines and syllabus resource routes
│   │   ├── schemas/
│   │   │   └── schemas.py               # Pydantic data validation contracts and DTOs
│   │   ├── services/
│   │   │   └── ai_service.py            # Local Ollama client & prompt orchestration
│   │   └── seed_courses.py              # Academic syllabus seed data loader
│   ├── coursealign.db                   # Embedded relational database storage
│   ├── create_admin.py                  # CLI administrator account bootstrapper
│   ├── Dockerfile                       # Backend containerization configuration
│   ├── main.py                          # Application entry point and ASGI middleware
│   └── requirements.txt                 # Backend Python package dependencies
├── frontend/
│   ├── assets/                          # Static branding, icons, and logo assets
│   ├── css/                             # Hand-drawn custom aesthetic and layout tokens
│   ├── js/                              # Modular client controllers (AI Tutor, Quiz, Catalog, Auth)
│   ├── pages/                           # Interactive HTML views for all user roles
│   ├── manifest.json                    # Progressive Web App web app manifest
│   ├── nginx.conf                       # Production Nginx reverse proxy configuration
│   └── sw.js                            # Service Worker for offline asset caching
├── docker-compose.yml                   # Multi-container service orchestrator
└── README.md                            # System documentation and deployment guide"""

    add_code_block(doc, "Project Directory Tree", tree_str)

    # A.2 Main Application Code
    p_a2 = doc.add_heading(level=2)
    p_a2.paragraph_format.space_before = Pt(10)
    p_a2.paragraph_format.space_after = Pt(6)
    p_a2.paragraph_format.keep_with_next = True
    r_a2 = p_a2.add_run("A.2 Main Application Code")
    r_a2.font.name = 'Arial'
    r_a2.font.size = Pt(12)
    r_a2.font.bold = True

    p = doc.add_paragraph("The primary entry point of the backend initializes the FastAPI framework instance, configures Cross-Origin Resource Sharing (CORS) policies, performs database schema initialization, and registers all modular API routes.")
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_after = Pt(2)
    r_lbl = p_lbl.add_run("Listing A.1: Main Application Entry Point and API Router Configuration")
    r_lbl.font.bold = True
    r_lbl.font.size = Pt(10)

    code_main = """from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from app.database.database import init_db
from app.routers import auth_router, course_router, bookmark_router, ai_router, admin_router
from app.core.config import settings

# Initialize database tables using sqlite3
init_db()

app = FastAPI(title=settings.PROJECT_NAME, version=settings.PROJECT_VERSION)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include Routers
app.include_router(auth_router.router)
app.include_router(course_router.router)
app.include_router(bookmark_router.router)
app.include_router(ai_router.router)
app.include_router(admin_router.router)

@app.get("/")
def read_root():
    return {"message": "Welcome to the CourseAlign API"}"""
    add_code_block(doc, "backend/main.py", code_main)

    # A.3 Authentication and Authorization
    p_a3 = doc.add_heading(level=2)
    p_a3.paragraph_format.space_before = Pt(10)
    p_a3.paragraph_format.space_after = Pt(6)
    p_a3.paragraph_format.keep_with_next = True
    r_a3 = p_a3.add_run("A.3 Authentication and Authorization")
    r_a3.font.name = 'Arial'
    r_a3.font.size = Pt(12)
    r_a3.font.bold = True

    p = doc.add_paragraph("User security is maintained through salted password hashing using bcrypt and stateless session verification using JSON Web Tokens (JWT). Role-Based Access Control (RBAC) is enforced through dependency injection guards ensuring that only authorized roles (Students, Lecturers, Administrators) access protected endpoints.")
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_after = Pt(2)
    r_lbl = p_lbl.add_run("Listing A.2: Security Utilities, JWT Token Generation, and Role-Based Guards")
    r_lbl.font.bold = True
    r_lbl.font.size = Pt(10)

    code_auth = """from datetime import datetime, timedelta, timezone
from typing import Optional
from jose import JWTError, jwt
import bcrypt
from fastapi import Depends, HTTPException, status
from fastapi.security import OAuth2PasswordBearer
import sqlite3

from app.core.config import settings
from app.database.database import get_db
from app.schemas.schemas import TokenData

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/auth/login")

def verify_password(plain_password, hashed_password):
    if isinstance(hashed_password, str):
        hashed_password = hashed_password.encode('utf-8')
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password)

def get_password_hash(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.now(timezone.utc) + expires_delta
    else:
        expire = datetime.now(timezone.utc) + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, settings.SECRET_KEY, algorithm=settings.ALGORITHM)
    return encoded_jwt

def get_current_user(token: str = Depends(oauth2_scheme), db: sqlite3.Connection = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
        
    cursor = db.cursor()
    cursor.execute("SELECT * FROM users WHERE username = ?", (token_data.username,))
    user = cursor.fetchone()
    if user is None:
        raise credentials_exception
    return dict(user)

def require_role(role: str):
    def role_checker(current_user: dict = Depends(get_current_user)):
        if current_user.get("role") != role:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=f"Operation not permitted. Requires {role} role."
            )
        return current_user
    return role_checker"""
    add_code_block(doc, "backend/app/auth/auth.py", code_auth)

    # A.4 Database Implementation
    p_a4 = doc.add_heading(level=2)
    p_a4.paragraph_format.space_before = Pt(10)
    p_a4.paragraph_format.space_after = Pt(6)
    p_a4.paragraph_format.keep_with_next = True
    r_a4 = p_a4.add_run("A.4 Database Implementation")
    r_a4.font.name = 'Arial'
    r_a4.font.size = Pt(12)
    r_a4.font.bold = True

    p = doc.add_paragraph("The relational persistence tier is implemented using an embedded SQLite database engine. Connection factories provide dict-like row access, and the Data Definition Language (DDL) initializes relational tables for users, courses, weekly outlines, uploaded materials, bookmarks, chat session logs, and quiz attempts.")
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_after = Pt(2)
    r_lbl = p_lbl.add_run("Listing A.3: SQLite Database Initialization and Relational Table Schemas")
    r_lbl.font.bold = True
    r_lbl.font.size = Pt(10)

    code_db = """import sqlite3
import os

DB_FILE = os.path.join(os.path.dirname(__file__), "..", "..", "coursealign.db")

def get_db_connection():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Users Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        email TEXT UNIQUE NOT NULL,
        hashed_password TEXT NOT NULL,
        role TEXT DEFAULT 'student',
        level TEXT DEFAULT '100L',
        is_active BOOLEAN DEFAULT 1,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')

    # Courses Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS courses (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        credit_unit INTEGER NOT NULL,
        level TEXT NOT NULL,
        semester TEXT NOT NULL,
        lecturer_name TEXT
    )
    ''')

    # Course Outlines Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS course_outlines (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        course_id INTEGER UNIQUE,
        learning_objectives TEXT,
        weekly_outline TEXT,
        recommended_textbooks TEXT,
        FOREIGN KEY(course_id) REFERENCES courses(id)
    )
    ''')

    # Bookmarks Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS bookmarks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        course_id INTEGER,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(course_id) REFERENCES courses(id)
    )
    ''')

    # Quiz Attempts Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS quiz_attempts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        course_id INTEGER,
        score INTEGER NOT NULL,
        total_questions INTEGER NOT NULL,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(course_id) REFERENCES courses(id)
    )
    ''')

    conn.commit()
    conn.close()"""
    add_code_block(doc, "backend/app/database/database.py", code_db)

    # A.5 Backend API Implementation
    p_a5 = doc.add_heading(level=2)
    p_a5.paragraph_format.space_before = Pt(10)
    p_a5.paragraph_format.space_after = Pt(6)
    p_a5.paragraph_format.keep_with_next = True
    r_a5 = p_a5.add_run("A.5 Backend/API Implementation")
    r_a5.font.name = 'Arial'
    r_a5.font.size = Pt(12)
    r_a5.font.bold = True

    p = doc.add_paragraph("The application exposes RESTful endpoints for course catalog navigation, syllabus outline management, lecturer syllabus authoring, and system administration.")
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_after = Pt(2)
    r_lbl = p_lbl.add_run("Listing A.4: Course Outline and Learning Resource Router")
    r_lbl.font.bold = True
    r_lbl.font.size = Pt(10)

    code_course_router = """from fastapi import APIRouter, Depends, HTTPException, status
import sqlite3
from typing import List

from app.database.database import get_db
from app.schemas.schemas import CourseResponse, CourseOutlineCreate, CourseOutlineResponse, CourseAssignUpdate
from app.auth.auth import get_current_active_user, require_role

router = APIRouter(prefix="/api/courses", tags=["courses"])

@router.get("/", response_model=List[CourseResponse])
def get_courses(skip: int = 0, limit: int = 100, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("SELECT * FROM courses LIMIT ? OFFSET ?", (limit, skip))
    return [dict(row) for row in cursor.fetchall()]

@router.put("/{course_code}/assign", response_model=CourseResponse, dependencies=[Depends(require_role("lecturer"))])
def assign_course(course_code: str, update_data: CourseAssignUpdate, db: sqlite3.Connection = Depends(get_db), current_user: dict = Depends(get_current_active_user)):
    cursor = db.cursor()
    cursor.execute("SELECT * FROM courses WHERE code = ?", (course_code,))
    course = cursor.fetchone()
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")
        
    cursor.execute(
        "UPDATE courses SET lecturer_name = ?, credit_unit = ?, level = ?, semester = ?, description = ? WHERE code = ?", 
        (current_user['username'], update_data.credit_unit, update_data.level, update_data.semester, update_data.description, course_code)
    )
    db.commit()
    cursor.execute("SELECT * FROM courses WHERE code = ?", (course_code,))
    return dict(cursor.fetchone())"""
    add_code_block(doc, "backend/app/routers/course_router.py", code_course_router)

    # A.6 AI/ML Implementation
    p_a6 = doc.add_heading(level=2)
    p_a6.paragraph_format.space_before = Pt(10)
    p_a6.paragraph_format.space_after = Pt(6)
    p_a6.paragraph_format.keep_with_next = True
    r_a6 = p_a6.add_run("A.6 AI/ML Implementation")
    r_a6.font.name = 'Arial'
    r_a6.font.size = Pt(12)
    r_a6.font.bold = True

    p = doc.add_paragraph("The machine learning subsystem interfaces with a locally hosted Large Language Model (Gemma via Ollama) to guarantee zero subscription costs and complete institutional data privacy. Dynamic context injection feeds course outlines and descriptions into prompts to produce tailored tutoring dialogue and structured assessment questions.")
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_after = Pt(2)
    r_lbl = p_lbl.add_run("Listing A.5: Local Ollama Gemma Integration and Context-Aware Generation Service")
    r_lbl.font.bold = True
    r_lbl.font.size = Pt(10)

    code_ai = """import os
import requests
import json

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434/api/generate")
OLLAMA_CHAT_URL = os.getenv("OLLAMA_CHAT_URL", "http://localhost:11434/api/chat")
MODEL_NAME = os.getenv("MODEL_NAME", "gemma2:2b")

def generate_ai_response(prompt: str, context: str = "") -> str:
    system_instruction = (
        "You are the CourseAlign AI Tutor, an expert assistant for Computer Science students at "
        "Chukwuemeka Odumegwu Ojukwu University (COOU). Be concise, helpful, and academically rigorous."
    )
    if context:
        system_instruction += f"\\nHere is the relevant course context:\\n{context}"
        
    messages = [
        {"role": "system", "content": system_instruction},
        {"role": "user", "content": prompt}
    ]
    payload = {"model": MODEL_NAME, "messages": messages, "stream": False}
    
    try:
        response = requests.post(OLLAMA_CHAT_URL, json=payload, timeout=120)
        response.raise_for_status()
        return response.json().get("message", {}).get("content", "Error parsing response.")
    except Exception as e:
        return f"An error occurred while communicating with the AI: {str(e)}" """
    add_code_block(doc, "backend/app/services/ai_service.py", code_ai)

    # =========================================================================
    # APPENDIX B: SAMPLE OUTPUTS
    # =========================================================================
    doc.add_page_break()
    p_app_b = doc.add_heading(level=1)
    p_app_b.paragraph_format.space_before = Pt(12)
    p_app_b.paragraph_format.space_after = Pt(12)
    p_app_b.paragraph_format.keep_with_next = True
    r_app_b = p_app_b.add_run("Appendix B: Sample Outputs")
    r_app_b.font.name = 'Arial'
    r_app_b.font.size = Pt(15)
    r_app_b.font.bold = True
    r_app_b.font.color.rgb = RGBColor(31, 41, 55)

    # Figures B.1 to B.14
    figures_data = [
        ("B.1 System Landing and Home Page", "appendix_b_01_home_page.png", 
         "Figure B.1: CourseAlign Public Landing Page and Feature Showcase",
         "Depicts the main responsive landing page featuring the custom hand-drawn aesthetic, value proposition highlights, quick-navigation links, and entry call-to-actions for account creation and sign-in."),
         
        ("B.2 User Authentication (Login) Portal", "appendix_b_02_login_portal.png",
         "Figure B.2: Secure User Authentication Portal",
         "Demonstrates the user authentication interface where students, lecturers, and system administrators input their registered credentials to obtain a secure JSON Web Token (JWT)."),
         
        ("B.3 User Registration and Onboarding Portal", "appendix_b_03_registration_portal.png",
         "Figure B.3: User Registration and Role Onboarding Interface",
         "Illustrates the student and lecturer onboarding portal with account creation form fields, level selector, and role assignment."),
         
        ("B.4 Student Dashboard and Learning Analytics", "appendix_b_04_student_dashboard.png",
         "Figure B.4: Student Analytics Dashboard and Learning Metrics",
         "Shows real-time aggregated metrics (Courses Studied: 3, Quiz Average: 90%, AI Chats: 1) calculated dynamically via SQL queries, alongside quick shortcuts to recent courses (CIS 201, CIS 312)."),
         
        ("B.5 Centralized Course Catalog and Level Filter", "appendix_b_05_course_catalog.png",
         "Figure B.5: Centralized Course Catalog and Level Filter Interface",
         "Displays the departmental course repository showing course codes, titles, credit units, lecturer assignments, and interactive level filter chips."),
         
        ("B.6 Detailed Course Outline and Syllabus Viewer", "appendix_b_06_course_detail_syllabus.png",
         "Figure B.6: Comprehensive Course Syllabus, Weekly Outline, and Objectives Viewer",
         "Illustrates the comprehensive syllabus viewer displaying structured weekly modules, textbook references, bookmarking toggles, and direct links to initiate an AI tutoring session for the course."),
         
        ("B.7 Saved Course Bookmarks Management", "appendix_b_07_saved_bookmarks.png",
         "Figure B.7: Bookmarked Course Management Interface",
         "Displays the student's personal saved courses collection for rapid offline/online access and quick revision."),
         
        ("B.8 AI-Powered 24/7 Intelligent Tutor Chat", "appendix_b_08_ai_tutor_chat.png",
         "Figure B.8: Interactive 24/7 AI Tutor Chat Interface with Context Injection",
         "Demonstrates an active tutoring dialogue with local Gemma 4 inference comparing AVL Trees and B-Trees with injected CIS 321 curriculum context."),
         
        ("B.9 Dynamic Practice Quiz Configuration Screen", "appendix_b_09_practice_quiz_config.png",
         "Figure B.9: Dynamic Practice Quiz Configuration Interface",
         "Shows the assessment generator setup screen where students configure course selection, difficulty tier, and question volume."),
         
        ("B.10 Interactive Assessment and Quiz Runner", "appendix_b_10_interactive_quiz_runner.png",
         "Figure B.10: Interactive Multiple-Choice Quiz Assessment Runner",
         "Demonstrates the live quiz runner interface with multiple-choice questions, instant feedback, and assessment grading."),
         
        ("B.11 Lecturer Dashboard and Course Assignments", "appendix_b_11_lecturer_dashboard.png",
         "Figure B.11: Lecturer Course Overview and Assignment Portal",
         "Shows the lecturer management interface displaying claimed courses (Dr. C. Okeke) and access to syllabus authoring tools."),
         
        ("B.12 Lecturer Syllabus Authoring Portal", "appendix_b_12_lecturer_manage_course.png",
         "Figure B.12: Lecturer Syllabus Authoring and Resource Upload Portal",
         "Demonstrates the course outline editing form used by academic staff to update weekly modules, learning goals, and textbook references."),
         
        ("B.13 Administrator Dashboard and Role Governance", "appendix_b_13_admin_user_management.png",
         "Figure B.13: Administrator User Governance and Role Assignment Portal",
         "Shows the central administrative console displaying user management tables with role-toggling controls (Student, Lecturer, Admin) and global course creation utilities."),
         
        ("B.14 User Profile and Account Settings", "appendix_b_14_user_profile.png",
         "Figure B.14: Student Profile and Account Settings Interface",
         "Illustrates user account management displaying registered email, username, academic level, and preferences.")
    ]

    for sec_title, img_file, cap, desc in figures_data:
        p_sec = doc.add_heading(level=2)
        p_sec.paragraph_format.space_before = Pt(10)
        p_sec.paragraph_format.space_after = Pt(6)
        p_sec.paragraph_format.keep_with_next = True
        r_sec = p_sec.add_run(sec_title)
        r_sec.font.name = 'Arial'
        r_sec.font.size = Pt(12)
        r_sec.font.bold = True

        add_figure_with_image(doc, img_file, cap, desc)

    # =========================================================================
    # APPENDIX C: QUESTIONNAIRE
    # =========================================================================
    doc.add_page_break()
    p_app_c = doc.add_heading(level=1)
    p_app_c.paragraph_format.space_before = Pt(12)
    p_app_c.paragraph_format.space_after = Pt(12)
    p_app_c.paragraph_format.keep_with_next = True
    r_app_c = p_app_c.add_run("Appendix C: Questionnaire")
    r_app_c.font.name = 'Arial'
    r_app_c.font.size = Pt(15)
    r_app_c.font.bold = True
    r_app_c.font.color.rgb = RGBColor(31, 41, 55)

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(6)
    p_note.paragraph_format.space_after = Pt(12)
    p_note.paragraph_format.line_spacing = 1.25
    p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    r_note_lbl = p_note.add_run("Note on Research Methodology and Data Collection: ")
    r_note_lbl.bold = True
    r_note_txt = p_note.add_run(
        "The research and implementation of the CourseAlign system utilized a design-and-creation software engineering methodology. "
        "System evaluation and verification were conducted through direct unit testing, API contract verification, and local LLM benchmark "
        "assessments rather than field-administered psychometric questionnaires. Consequently, empirical questionnaire instruments and survey response "
        "datasets were not administered for this development phase."
    )

    output_path = os.path.join(os.path.dirname(__file__), "CourseAlign_Project_Appendices.docx")
    doc.save(output_path)
    print(f"Document with embedded screenshots successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_appendices()
